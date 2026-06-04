# -*- coding: utf-8 -*-
"""
브리핑 한 항목을 공공기관 개조식 워드(.docx)로 만든 뒤
용량을 최소화(4-part docx)해서 base64 한 줄로 stdout 에 출력한다.

사용법:
  python make_docx_b64.py <archive.js> [as_of]
    as_of 생략 시 배열 첫 항목(=최신).
출력: stdout 에 base64 문자열 1줄 (그 외 로그 일절 없음).
실패 시 비-0 종료코드 + stderr 메시지.

의존: python-docx.
"""
import sys, re, json, io, base64, zipfile

try:
    from docx import Document
    from docx.shared import Pt, Twips
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    sys.stderr.write("python-docx not installed. Run: python -m pip install python-docx --quiet --user\n")
    sys.exit(2)

FONT = "바탕체"
LINE = 1.3


def load_briefings(js_path):
    with open(js_path, "r", encoding="utf-8") as f:
        text = f.read()
    m = re.search(r"=\s*(\[.*\])\s*;?\s*$", text, re.S)
    if not m:
        raise ValueError("briefings array not found in " + js_path)
    return json.loads(m.group(1))


def report_from_briefing(b):
    rep = b.get("report")
    if rep and rep.get("sections"):
        return {"headline": rep.get("headline") or b.get("title") or "",
                "sections": rep["sections"]}
    headline = b.get("title") or ""
    sections, cur = [], None
    skip = {"안녕하십니까", "감사합니다.", "감사합니다",
            "금일 아시아 증시 시황 보고 드립니다."}
    for p in (b.get("paragraphs") or []):
        t = ("" if p is None else str(p)).strip()
        if not t or t in skip:
            continue
        if t[0] == "*":
            cur = {"head": "주요 지수 동향",
                   "points": [{"text": re.sub(r"^\*\s*", "", t), "subs": []}]}
            sections.append(cur)
        elif t[0] == "※":
            txt = re.sub(r"^※\s*", "", t)
            if cur and cur.get("points"):
                cur["points"][-1].setdefault("subs", []).append(txt)
            else:
                cur = {"head": txt, "points": []}
                sections.append(cur)
        else:
            cur = {"head": t, "points": []}
            sections.append(cur)
    return {"headline": headline, "sections": sections}


def set_east_asia(run):
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for k in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(k), FONT)


def set_word_wrap(par):
    pPr = par._p.get_or_add_pPr()
    ww = OxmlElement("w:wordWrap")
    ww.set(qn("w:val"), "0")
    pPr.insert_element_before(
        ww, "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE",
        "w:autoSpaceDN", "w:bidi", "w:adjustRightInd", "w:snapToGrid",
        "w:spacing", "w:ind", "w:contextualSpacing", "w:mirrorIndents",
        "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment",
        "w:textboxTightWrap", "w:outlineLvl", "w:divId", "w:cnfStyle",
        "w:rPr", "w:sectPr", "w:pPrChange")


def add_para(doc, text, *, size_pt, bold=False, underline=False,
             align=None, before=0, after=0,
             left=None, hanging=None, word_wrap=False):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if left is not None:
        pf.left_indent = Twips(left)
    if hanging is not None:
        pf.first_line_indent = Twips(-hanging)
    if word_wrap:
        set_word_wrap(p)
    r = p.add_run(text)
    r.font.size = Pt(size_pt)
    r.font.bold = bold
    r.font.underline = underline
    set_east_asia(r)
    return p


def build_docx_bytes(brief):
    rep = report_from_briefing(brief)
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(15)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for k in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(k), FONT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = LINE
    for sec in doc.sections:
        sec.top_margin = Twips(1134)
        sec.bottom_margin = Twips(1134)
        sec.left_margin = Twips(1134)
        sec.right_margin = Twips(1134)
    headline = rep.get("headline") or brief.get("title") or ""
    add_para(doc, headline, size_pt=22, bold=True, underline=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, after=18)
    if brief.get("as_of"):
        add_para(doc, "기준일 " + brief["as_of"], size_pt=12,
                 align=WD_ALIGN_PARAGRAPH.RIGHT, after=12)
    for sec in (rep.get("sections") or []):
        add_para(doc, "□ " + sec.get("head", ""), size_pt=15, bold=True,
                 before=12, after=12, left=450, hanging=450)
        for pt in (sec.get("points") or []):
            add_para(doc, "- " + pt.get("text", ""), size_pt=15,
                     before=6, after=6, left=600, hanging=300, word_wrap=True)
            for s in (pt.get("subs") or []):
                add_para(doc, "· " + s, size_pt=15,
                         before=3, after=3, left=750, hanging=300, word_wrap=True)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# 4-part minimal docx: 모든 서식이 인라인이므로 styles.xml/theme/numbering/customXml/docProps 모두 제거.
MIN_CONTENT_TYPES = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/></Types>'''

MIN_RELS = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>'''

MIN_DOC_RELS = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"></Relationships>'''


def minify_docx(full_bytes):
    """python-docx 출력에서 word/document.xml 만 꺼내 최소 zip 으로 다시 묶는다."""
    with zipfile.ZipFile(io.BytesIO(full_bytes), "r") as zin:
        doc_xml = zin.read("word/document.xml")
    out = io.BytesIO()
    with zipfile.ZipFile(out, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as zout:
        zout.writestr("[Content_Types].xml", MIN_CONTENT_TYPES)
        zout.writestr("_rels/.rels", MIN_RELS)
        zout.writestr("word/_rels/document.xml.rels", MIN_DOC_RELS)
        zout.writestr("word/document.xml", doc_xml)
    return out.getvalue()


def main():
    if len(sys.argv) < 2:
        sys.stderr.write("usage: make_docx_b64.py <archive.js> [as_of]\n")
        sys.exit(1)
    js_path = sys.argv[1]
    as_of = sys.argv[2] if len(sys.argv) > 2 else None
    arr = load_briefings(js_path)
    if not arr:
        sys.stderr.write("empty briefings archive: " + js_path + "\n")
        sys.exit(1)
    brief = arr[0]
    if as_of:
        for b in arr:
            if b.get("as_of") == as_of:
                brief = b
                break
    full = build_docx_bytes(brief)
    small = minify_docx(full)
    b64 = base64.b64encode(small).decode("ascii")
    sys.stdout.write(b64)


if __name__ == "__main__":
    main()
